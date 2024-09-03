import customtkinter as ctk
import tkinter as tk
from mysql.connector import Error
import mysql.connector
from tkinter import messagebox
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

data_atual = datetime.date.today()

class GenereteStatistics(ctk.CTkToplevel):
    def __init__(self) -> None:
        super().__init__()        

        center_window(self, 350, 250)
        self.title("GERAR ESTATISTICA MENSAL/ANUAL")
        self.resizable(width=False, height=False)

        self.meses = {'01':'JANEIRO', '02':'FEVEREIRO', '03':'MARÇO', '04':'ABRIL', '05':'MAIO', '06':'JUNHO', '07':'JULHO', '08':'AGOSTO', '09':'SETEMBRO', 10:'OUTUBRO', 11:'NOVEMBRO', 12:'DEZEMBRO'}
        self.ano = ['2024', '2025', '2026', '2027', '2028']

        self.label_tipo = ctk.CTkLabel(self, text="Selecione o mês e ano que deseja emitir o relatório", font=("Helvetica", 15)).pack(padx=1, pady=10)
        self.mes = ctk.CTkComboBox(self, values=list(self.meses.values()), width=250)
        self.mes.pack(padx=10, pady=1)
        self.ano = ctk.CTkComboBox(self, values=self.ano, width=250)
        self.ano.pack(padx=10, pady=5)

        self.button_generate= ctk.CTkButton(self, text="Gerar Estatistica Mensal", command=lambda: self.create_docx_file(False), corner_radius=32, hover_color="#000000", border_width=1, width=200).pack(pady=10)
        self.button_generate_anual= ctk.CTkButton(self, text="Gerar Estatistica Anual", command=lambda: self.create_docx_file(True), corner_radius=32, hover_color="#000000", border_width=1, width=200)
        self.button_generate_anual.pack(pady=10)

    def get_lista_atualizada(self, anual):
        connection = create_connection("localhost", "root", "root", "cejusc")  

        list_a_realizar = []
        
        if int(self.ano.get()) == int(data_atual.year):
            mes_atual = '%02d' % data_atual.month
            ano_atual = int(data_atual.year)
        else:
            ano_atual = int(self.ano.get())
            mes_atual = 12

        if anual:
            total_sessoes_realizar1,  total_sessoes_realizar2 = get_sessoes_realizar(mes_atual, ano_atual, connection)
        else:
            total_sessoes_realizar1,  total_sessoes_realizar2 = get_sessoes_realizar(int(list(self.meses.keys())[list(self.meses.values()).index(self.mes.get())]), ano_atual, connection)

        #[(5, 6, 15, 16, 11, 31)]
            
        #[('Pré-Civil', Decimal('30'), Decimal('3'), Decimal('5')), ('Pré-Familia', Decimal('30'), Decimal('5'), Decimal('5')), ('Pro-Civil', Decimal('50'), Decimal('7'), Decimal('5')), ('Pro-Familia', Decimal('50'), Decimal('8'), Decimal('4'))]
            
        print(total_sessoes_realizar2)
        for i in range(4):
            list_a_realizar.append([total_sessoes_realizar2[i][0], (int(total_sessoes_realizar1[0][i]) + int(total_sessoes_realizar2[i][1]) - int(total_sessoes_realizar2[i][2]) - int(total_sessoes_realizar2[i][3]))])

        list_realizar2 = list_a_realizar.copy()

        list_a_realizar.append(['Total-Pré', list_realizar2[0][1] + list_realizar2[1][1]])
        list_a_realizar.append(['Total-Pro', list_a_realizar[2][1] + list_a_realizar[3][1]])
        return list_a_realizar, mes_atual
    
    def create_docx_file(self, anual):
        connection = create_connection("localhost", "root", "root", "cejusc")

        list_realizar, mes_atual = self.get_lista_atualizada(anual)
        atualiza_saldo(list_realizar, int(list(self.meses.keys())[list(self.meses.values()).index(self.mes.get())]), self.ano.get())

        document = Document()
        title = document.add_paragraph()
        doc_font_color = RGBColor(0, 0, 0)        

        if anual:        
            run1 = title.add_run(f"ESTATÍSTICA PRÉ-PROCESSUAL DO ANO DE {self.ano.get()}")
            doc_name = f"Estatistica {self.ano.get()}.docx"
        else:
            run1 = title.add_run(f"ESTATÍSTICA PRÉ-PROCESSUAL DO MÊS DE {self.mes.get()} DE {self.ano.get()}")
            doc_name = f"Estatistica {self.mes.get()}-{self.ano.get()}.docx"
        
        run1.font.color.rgb = doc_font_color
        run1.font.size = Pt(15)
        run1.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER      
        
        table = document.add_table(rows=16, cols=4)
        set_table_borders(table)
        header = table.rows[0].cells
        header[0].text = 'ASSUNTO'
        header[1].text = 'CÍVEL'
        header[2].text = 'FAMÍLIA'
        header[3].text = 'TOTAL'

        list_subject = ['Total acordo obtidos',  'Total sessões infrutíferas', 'Total audiências realizadas', 'Total audiências designadas', 'Total ausência do requerente', 'Total ausência do requerido', 'Total ausência das partes', 'Total sessões canceladas', 'Total sessões redesignadas', 'Total sessões não realizadas', 'Total sessões à realizar', 'Pauta em dias', 'Total JG Dativo', 'Total JG Adv', 'Total de sessões gratuitas']

        counter_subject = 1
        for subject in list_subject:
            rows = table.rows[counter_subject].cells
            rows[0].text = subject
            counter_subject = counter_subject + 1

        list_tipos_pre = [['Pré-Civil', 'coalesce(sum(pauta_dias),0)', f'{list_realizar[0][1]}'], ['Pré-Familia', 'coalesce(sum(pauta_dias),0)', f'{list_realizar[1][1]}'], ["""Pré-Civil', 'Pré-Familia""", """'-'""", f'{list_realizar[4][1]}']]

        for i in range(3):
            numeros_pre_civil = self.get_numbers(list_tipos_pre[i], anual, connection, mes_atual)
            counter_pre = 1
            for item in numeros_pre_civil:
                rows = table.rows[counter_pre].cells
                rows[i + 1].text = str(item)
                counter_pre = counter_pre + 1

        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = doc_font_color
                        run.font.size = Pt(11)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        document.add_page_break()
        
        title_page2 = document.add_paragraph()
        if anual:        
            run2 = title_page2.add_run(f"ESTATÍSTICA PROCESSUAL DO ANO DE {self.ano.get()}")
        else:
            run2 = title_page2.add_run(f"ESTATÍSTICA PROCESSUAL DO MÊS DE {self.mes.get()} DE {self.ano.get()}")

        run2.font.color.rgb = doc_font_color
        run2.font.size = Pt(15)
        run2.font.bold = True
        title_page2.alignment = WD_ALIGN_PARAGRAPH.CENTER


        table = document.add_table(rows=16, cols=4)
        #table.style = 'Medium Grid 2'
        header = table.rows[0].cells
        header[0].text = 'ASSUNTO'
        header[1].text = 'CÍVEL'
        header[2].text = 'FAMÍLIA'
        header[3].text = 'TOTAL'
        set_table_borders(table)

        counter_subject_table2 = 1
        for subject in list_subject:
            rows = table.rows[counter_subject_table2].cells
            rows[0].text = subject
            counter_subject_table2 = counter_subject_table2 + 1

        list_tipos_pro = [['Pro-Civil', 'coalesce(sum(pauta_dias),0)', f'{list_realizar[2][1]}'], ['Pro-Familia', 'coalesce(sum(pauta_dias),0)', f'{list_realizar[3][1]}'], ["""Pro-Civil', 'Pro-Familia""", """'-'""", f'{list_realizar[5][1]}']]

        for t in range(3):
            numeros_pro_civil = self.get_numbers(list_tipos_pro[t], anual, connection, mes_atual)
            counter_pro_civil = 1
            for item in numeros_pro_civil:
                rows = table.rows[counter_pro_civil].cells 
                rows[t + 1].text = str(item)
                counter_pro_civil = counter_pro_civil + 1

        # Formata a primeira coluna em negrito
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = doc_font_color
                        run.font.size = Pt(11)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        

        #document.save(doc_name)
        document.save("C:/Users/CEJUSC-01/Desktop/" + doc_name)
        messagebox.showinfo("Sucesso","Estatística exportada com sucesso")



    def get_numbers(self, tipo, anual, connection, mes_atual):
        mes = list(self.meses.keys())[list(self.meses.values()).index(self.mes.get())]
        ano = self.ano.get()
        
        query_mes = f"""
                    select 
                    coalesce(sum(total_acordos_obtidos),0), 
                    coalesce(sum(total_sessoes_infrutiferas),0),
                    coalesce(sum(total_audiencias_realizadas),0),
                    coalesce(sum(total_audiencias_designadas),0),
                    coalesce(sum(total_ausencia_requerente),0),
                    coalesce(sum(total_ausencia_requerido),0),
                    coalesce(sum(total_ausencia_partes),0),
                    coalesce(sum(total_sessoes_canceladas),0),
                    coalesce(sum(total_sessoes_redesignadas),0),
                    coalesce(sum(Total_sessoes_nao_realizadas),0),
                    '{tipo[2]}',
                    {tipo[1]},
                    coalesce(sum(total_jg_Dativo),0),
                    coalesce(sum(total_jg_adv),0),
                    coalesce(sum(Total_sessões_gratuitas),0)
                    from estatistica_mensal 
                    where 
                    month(data) = '{mes}' and year(data) = '{ano}'
                    and tipo_processo in ('{tipo[0]}')
                    """

        query_anual = f"""
                    select 
                    sum(total_acordos_obtidos) as total_acordos_obtidos, 
                    sum(total_sessoes_infrutiferas),
                    sum(total_audiencias_realizadas),
                    sum(total_audiencias_designadas),
                    sum(total_ausencia_requerente),
                    sum(total_ausencia_requerido),
                    sum(total_ausencia_partes),
                    sum(total_sessoes_canceladas),
                    sum(total_sessoes_redesignadas),
                    sum(Total_sessoes_nao_realizadas),
                    '{tipo[2]}',
                    {tipo[1]},
                    sum(total_jg_Dativo),
                    sum(total_jg_adv),
                    sum(Total_sessões_gratuitas)
                    from estatistica_mensal 
                    where 
                    year(data) = '{ano}'
                    and month(data) = '{mes_atual}'
                    and tipo_processo in ('{tipo[0]}')
                    """
        cursor = connection.cursor()
        try:
            if anual:
                cursor.execute(query_anual)
            else:
                cursor.execute(query_mes)
            myresult = cursor.fetchall()
            return myresult[0]            
        except Error as e:
            print(e)

def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Tamanho da borda
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')  # Cor clara (cinza claro)

        tblBorders.append(border)

    tblPr.append(tblBorders)


def create_connection(host_name, user_name, user_password, db_name):
    connection = None
    try:
        connection = mysql.connector.connect(
            host=host_name,
            user=user_name,
            passwd=user_password,
            database=db_name
        )
        print("Conexão com banco de dados bem-sucedida")
    except Error as e:
        messagebox.showerror("ERRO", f"ERRO AO CONECTAR COM BANCO DE DADOS: {e}")
    return connection

def center_window(window, width, height):
    window.update_idletasks()
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    window.geometry(f"{width}x{height}+{x}+{y}")


#Total audiencias a realizar é a soma do saldo do mes anterior + designadas no mes atual - audiencias realizadas no mes atual - audiencias nao realizadas no mes atual

def get_sessoes_realizar(month, ano_atual, conexao): 

    query_sessoes_1 = f"""
    select 
    Pre_civil,
    Pre_familia,
    Pro_civil,
    Pro_familia
    from saldo_sessoes_realizar em     
    where id = (SELECT MAX(ID) FROM saldo_sessoes_realizar where mes = {int(month) - 1})
    and ano = {ano_atual}
    """

    query_sessoes_2 = f"""
    SELECT 
    DISTINCT tipo_processo,
    (SELECT 
        COALESCE(SUM(total_audiencias_designadas), 0) 
    FROM 
        estatistica_mensal 
    WHERE 
        tipo_processo = em.tipo_processo AND 
        MONTH(data) = '{month}' AND 
        YEAR(data) = '{ano_atual}'),
    (SELECT 
        COALESCE(SUM(total_audiencias_realizadas), 0) 
    FROM 
        estatistica_mensal 
    WHERE 
        tipo_processo = em.tipo_processo AND 
        MONTH(data) = '{month}' AND 
        YEAR(data) = '{ano_atual}'),
    (SELECT 
        COALESCE(SUM(Total_sessoes_nao_realizadas), 0) 
    FROM 
        estatistica_mensal 
    WHERE 
        tipo_processo = em.tipo_processo AND 
        MONTH(data) = '{month}' AND 
        YEAR(data) = '{ano_atual}')
    FROM 
        estatistica_mensal em
    ORDER BY 
        tipo_processo;
    """

    cursor = conexao.cursor()
    try:       
        cursor.execute(query_sessoes_1)
        myresult_1 = cursor.fetchall() 

        cursor.execute(query_sessoes_2)
        myresult_2 = cursor.fetchall() 

        if len(query_sessoes_2) < 1:
            messagebox.showerror("ERRO", f"NÃO EXISTEM REGISTROS NO PERÍODO INFORMADO")
        else:
            return myresult_1, myresult_2
    except Error as e:
        print(e)


def atualiza_saldo(lista, mes, ano):
        connection = create_connection("localhost", "root", "root", "cejusc") 
        val = [mes, lista[1][1], lista[0][1], lista[3][1], lista[2][1], ano, lista[4][1], lista[5][1]]
        #val = [mes, lista[2][1], lista[3][1], lista[0][1], lista[1][1], ano, lista[4][1], lista[5][1]]
        query = f"""
        INSERT INTO saldo_sessoes_realizar (mes, Pre_familia, Pre_civil, Pro_familia, Pro_civil, ano, total_pre, total_pro)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"""
        
        cursor = connection.cursor()
        try:
            cursor.execute(query ,val)
            connection.commit()
            #messagebox.showinfo("Sucesso","Estatística inserida com sucesso")
        except Error as e:
            messagebox.showerror("ERRO", f"Erro ao inserir Estatística: {e}")